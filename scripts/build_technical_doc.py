"""Build the submission-ready FinGuard AI technical document."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs" / "FinGuard_AI_Technical_Documentation.docx"
SCREENSHOT = ROOT / "docs" / "assets" / "dashboard-explanation-en.png"
PARAMS = json.loads((ROOT / "models" / "params.json").read_text(encoding="utf-8"))

NAVY = RGBColor(10, 14, 26)
BLUE = RGBColor(74, 122, 255)
INK = RGBColor(38, 48, 73)
MUTED = RGBColor(100, 111, 137)
RED = RGBColor(196, 52, 63)
WHITE = RGBColor(255, 255, 255)
LIGHT = "F2F4F7"


def set_font(run, *, size=None, bold=None, color=None, name="Calibri", italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + side
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE if style_name != "Heading 3" else INK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def setup_section(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("FINGUARD AI  |  TECHNICAL DOCUMENTATION")
    set_font(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("FinGuard AI  ·  MVP v2.0  ·  Page ")
    set_font(run, size=8.5, color=MUTED)
    add_page_number(footer)


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc, text, *, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, bold=True, color=INK)
        text = text[len(bold_lead):]
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    for idx, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, LIGHT)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(header) < 16 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        set_font(run, size=9.5, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for idx, (value, width) in enumerate(zip(row, widths)):
            cells[idx].width = Inches(width)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == len(row) - 1 and len(str(value)) < 16 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            set_font(run, size=9.3, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_callout(doc, title, text, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    shade(cell, "EEF3FF")
    p = cell.paragraphs[0]
    run = p.add_run(title.upper() + "\n")
    set_font(run, size=9, bold=True, color=color)
    run = p.add_run(text)
    set_font(run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build():
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        setup_section(section)

    # Editorial-cover opening with the standard_business_brief style system.
    doc.add_paragraph().paragraph_format.space_after = Pt(70)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("EXPLAINABLE UPI FRAUD DETECTION")
    set_font(run, size=10, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("FinGuard AI")
    set_font(run, size=30, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("Technical Architecture, Model Card, API, Security, and Deployment")
    set_font(run, size=14, color=MUTED)
    add_callout(
        doc,
        "Prototype promise",
        "Detect suspicious UPI-like transactions, expose the top model factors, and turn them into a clear English or Hindi safety action.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run("MVP v2.0  |  20 June 2026  |  Submission package")
    set_font(run, size=10, bold=True, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "1. Executive summary")
    add_body(doc, "FinGuard AI converts a black-box fraud alert into an actionable explanation. A validated transaction is transformed into behavioral features, scored by an Isolation Forest, attributed with SHAP, and explained in English or Hindi. FastAPI provides the integration surface and Streamlit provides a user and analyst dashboard.")
    add_callout(doc, "Offline-first demo", "The included model, seeded synthetic feed, and deterministic bilingual explanation layer allow the complete judge flow to run without Groq or internet access.")

    add_heading(doc, "2. Goals and boundaries")
    add_heading(doc, "Prototype goals", 2)
    add_bullets(doc, [
        "Return a calibrated 0-100 risk score and top-three SHAP factors.",
        "Explain the same evidence in plain English or Hindi.",
        "Support live-feed simulation, manual case review, and seven-day analytics.",
        "Expose predictable API contracts that can sit beside an existing fraud stack.",
    ])
    add_heading(doc, "Explicit boundaries", 2)
    add_bullets(doc, [
        "No connection to NPCI, banks, real accounts, or real transaction streams.",
        "No automatic payment decline, reversal, reporting, or enforcement action.",
        "No UPI PIN, account number, phone number, or personal identifier is required.",
        "Synthetic holdout metrics are not production-bank performance claims.",
    ])

    add_heading(doc, "3. Architecture")
    add_table(doc, ["Layer", "Technology", "Responsibility"], [
        ("Experience", "Streamlit", "Seeded live feed, manual analysis, alerts, SHAP charts, and analytics."),
        ("API", "FastAPI + Pydantic", "Input validation, prediction, explanation, health, metrics, and demo endpoints."),
        ("Detection", "Isolation Forest", "Learns legitimate behavior and produces an anomaly normality score."),
        ("Explainability", "SHAP TreeExplainer", "Attributes transformed features and aggregates them to user-facing fields."),
        ("Language", "Groq / Llama or local fallback", "Produces constrained two-sentence English or Hindi safety guidance."),
        ("Packaging", "Docker + Render config", "Provides a stateless API deployment unit and health check."),
    ], [1.05, 1.55, 3.9])

    if SCREENSHOT.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(SCREENSHOT), width=Inches(6.35))
        caption = doc.add_paragraph("Figure 1. Verified Streamlit fraud alert with SHAP factors and offline explanation.")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)
        for run in caption.runs:
            set_font(run, size=8.5, italic=True, color=MUTED)

    add_heading(doc, "4. Data and feature engineering")
    add_body(doc, "The generated dataset contains 9,000 normal and 1,000 fraudulent UPI-like records. Fraud is injected through late-night high-value payments, device and location jumps, velocity attacks, and social-engineering payments to new merchants.")
    add_table(doc, ["Input feature", "Type", "Meaning"], [
        ("amount", "Float", "Payment amount in INR."),
        ("hour", "Integer", "Local hour from 0 to 23."),
        ("merchant_category", "Category", "Ten merchant groups; one-hot encoded."),
        ("device_change", "Binary", "Device differs from the established fingerprint."),
        ("geo_distance_km", "Float", "Distance from the prior transaction location."),
        ("velocity_per_hour", "Integer", "Transactions observed in the current hour."),
        ("is_new_merchant", "Binary", "Merchant is new to the payer."),
    ], [1.55, 0.85, 4.10])
    add_body(doc, "Transaction ID, label, and fraud-pattern metadata are not model inputs.")

    add_heading(doc, "5. Training and model card")
    add_bullets(doc, [
        "Create a stratified 80/20 split with random seed 42.",
        "Fit preprocessing and Isolation Forest only on normal training rows.",
        "Standardize numeric inputs and one-hot encode merchant category.",
        "Map normality to 0-100 risk with headroom for severe anomalies.",
        "Select the fraud-class F1-maximizing threshold on the held-out partition.",
        "Build and persist a SHAP TreeExplainer for the fitted forest.",
    ])
    add_table(doc, ["Held-out metric", "Result"], [
        ("Test rows", f"{PARAMS['test_rows']:,}"),
        ("Fraud rows", f"{PARAMS['test_fraud_rows']:,}"),
        ("Fraud precision", f"{PARAMS['precision_fraud']:.4f}"),
        ("Fraud recall", f"{PARAMS['recall_fraud']:.4f}"),
        ("Fraud F1", f"{PARAMS['f1']:.4f}"),
        ("ROC-AUC", f"{PARAMS['roc_auc']:.4f}"),
        ("Calibrated threshold", f"{PARAMS['threshold']:.2f}"),
    ], [4.7, 1.8])
    add_callout(doc, "Interpretation", "The high score is expected because the synthetic attack rules are deliberately separable. It proves the implementation path, not external validity. Production evaluation requires institution data and temporal validation.", color=RED)

    add_heading(doc, "6. Explainability and language safety")
    add_body(doc, "Numeric SHAP contributions map directly to their original input. One-hot merchant contributions are summed into a single merchant-category contribution. The API returns value, bilingual label, signed contribution, and risk direction for each factor.")
    add_bullets(doc, [
        "Only risk score and top contributing factors are sent to the LLM.",
        "The system prompt forbids certainty, unsupported reasons, and ML jargon.",
        "The dashboard escapes generated text before inserting it into HTML.",
        "Timeouts, missing keys, and malformed responses fall back locally without leaking errors.",
    ])

    add_heading(doc, "7. API contract")
    add_table(doc, ["Method", "Endpoint", "Purpose"], [
        ("GET", "/health", "Model readiness and explanation mode."),
        ("GET", "/model/metrics", "Persisted training and evaluation metadata."),
        ("POST", "/predict", "Risk, threshold, decision, version, and SHAP factors."),
        ("POST", "/explain", "Prediction plus English or Hindi explanation."),
        ("GET", "/demo/feed", "Mixed synthetic batch with simulated ground truth."),
    ], [0.8, 1.45, 4.25])
    add_body(doc, "OpenAPI documentation is generated automatically at /docs. Numeric inputs have explicit upper and lower bounds, and language is restricted to English or Hindi.")

    add_heading(doc, "8. Security, privacy, and responsible use")
    add_heading(doc, "Implemented in the MVP", 2)
    add_bullets(doc, [
        "Strict request validation, limited CORS methods, environment-based secrets, and no key logging.",
        "No personal or payment credential is needed for scoring or explanation.",
        "Graceful LLM failure and HTML escaping reduce availability and injection risk.",
        "The tool is decision support; it does not take consequential payment action.",
    ])
    add_heading(doc, "Required before production", 2)
    add_bullets(doc, [
        "OAuth2 or mTLS, role-based access, tenant isolation, rate limits, and immutable audit logs.",
        "TLS, managed secrets, encryption at rest, retention policy, and incident response.",
        "Model registry, approval and rollback, drift detection, fairness slices, and human review.",
        "Consent and privacy review before any WhatsApp or user-notification integration.",
    ])

    add_heading(doc, "9. Scalability and deployment")
    add_body(doc, "The API loads artifacts once and performs CPU-local inference, so it can scale horizontally behind a load balancer. At bank volume, scoring should consume transaction events from a queue, use a feature store for behavioral history, persist alerts separately, and invoke the LLM only for flagged transactions. Redis can enforce rate limits and cache repeated explanations.")
    add_body(doc, "The included Dockerfile serves FastAPI on the platform PORT. The Render definition and Streamlit configuration are included, but no hosted demo URL is part of this package.")

    add_heading(doc, "10. Verification and failure behavior")
    add_table(doc, ["Condition", "Expected behavior"], [
        ("Model missing", "API returns 503; dashboard shows setup instructions."),
        ("Groq unavailable", "Feature-grounded local English or Hindi explanation."),
        ("Invalid input", "Pydantic rejects the request with HTTP 422."),
        ("Unknown merchant", "Encoder ignores the unseen category; other signals continue."),
        ("Visible UI error", "Automated browser capture fails the QA run."),
    ], [2.05, 4.45])
    add_body(doc, "The automated suite currently contains eight passing tests and covers data rules, model risk ordering, missing fields, language grounding, and API validation. A smoke test confirms the end-to-end explain path.")

    add_heading(doc, "11. Limitations and next steps")
    add_bullets(doc, [
        "Replace synthetic data with de-identified, consented institution data and time-based splits.",
        "Compute device, location, merchant, and velocity deltas from trusted account history.",
        "Measure alert capacity, false-positive cost, calibration, regional slices, and drift.",
        "Add authenticated case management, acknowledgment, analyst notes, and audit export.",
        "Load-test scoring and explanation paths independently before a pilot.",
    ])

    add_heading(doc, "12. References")
    add_bullets(doc, [
        "NPCI: UPI Product Statistics - https://www.npci.org.in/what-we-do/upi/product-statistics",
        "Reserve Bank of India: Annual Report Publications - https://www.rbi.org.in/Scripts/AnnualReportPublications.aspx",
        "scikit-learn: IsolationForest documentation - https://scikit-learn.org/stable/modules/generated/sklearn.ensemble.IsolationForest.html",
        "SHAP documentation - https://shap.readthedocs.io/",
        "FastAPI documentation - https://fastapi.tiangolo.com/",
    ])
    add_body(doc, "Market figures supplied in the challenge brief should be refreshed from the latest NPCI and RBI releases immediately before final upload.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "FinGuard AI Technical Documentation"
    doc.core_properties.subject = "Explainable UPI Fraud Detection MVP"
    doc.core_properties.author = "FinGuard AI Team"
    doc.core_properties.keywords = "UPI, fraud detection, SHAP, Isolation Forest, FastAPI, Streamlit"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
